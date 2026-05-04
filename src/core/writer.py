"""WordWriter — fills a Word template with record data and saves the result.

Placeholder mapping (template → Excel column / computed value)
--------------------------------------------------------------
<Master Agreement>       → "Master Agreement"
<Supplier Name>          → "Supplier name"
<Effective Date>         → "Effective Date"  (formatted as Month DD, YYYY)
<Sub-Category>           → "Sub-Category"
<Platform>               → "Platform"
<Actual Payment>         → "Actual Payment"  (formatted as $X,XXX.XX)
<Capital Money Letter>   → "Actual Payment"  (uppercase English words)
<Signer>                 → "Signer"
<Title>                  → "Signer title"

Formatting rules applied after substitution
-------------------------------------------
* Body paragraph index 2 (0-based) → bold + centred
* Table cell containing "Total Settlement Payment (in USD)" → right-aligned bold
* Paragraphs containing "HP INC." or equal to "COMPANY" → left-aligned bold
* Entire document → Times New Roman 10 pt
"""
from __future__ import annotations

from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.config.paths import OUTPUT_DIR, TEMPLATE_ELAN_SYNA, TEMPLATE_GENERAL

_FONT_NAME = "Times New Roman"
_FONT_SIZE = Pt(10)

# Maps each Word placeholder to the corresponding record key.
# Sentinels (starting/ending with '__') are computed at runtime.
_PLACEHOLDER_MAP: dict[str, str] = {
    "<Master Agreement>": "Master Agreement",
    "<Supplier Name>": "Supplier name",
    "<Effective Date>": "__effective_date__",
    "<Sub-Category>": "Sub-Category",
    "<Platform>": "Platform",
    "<Actual Payment>": "__actual_payment__",
    "<Capital Money Letter>": "__capital_money__",
    "<Signer>": "Signer",
    "<Title>": "Signer title",
}


class WordWriter:
    """Fills a Word template with record data and saves the output file."""

    def __init__(self, output_dir: Path = OUTPUT_DIR) -> None:
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    def write(
        self,
        record: dict,
        is_elan_syna: bool,
        format_amount: Callable[[float], str],
        amount_to_words: Callable[[float], str],
        format_date: Callable[[object], str],
    ) -> Path:
        """Fill template, apply formatting, save, and return the output Path."""
        template_path = TEMPLATE_ELAN_SYNA if is_elan_syna else TEMPLATE_GENERAL
        doc = Document(str(template_path))

        replacements = _build_replacements(
            record, format_amount, amount_to_words, format_date
        )

        # --- Replace placeholders in body paragraphs ---
        for para in doc.paragraphs:
            _replace_in_paragraph(para, replacements)

        # --- Apply body-level formatting rules ---
        _apply_body_formatting(list(doc.paragraphs))

        # --- Replace + format inside tables ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, replacements)
                        _apply_table_cell_formatting(para)

        # --- Uniform font across the whole document ---
        _apply_global_font(doc)

        # --- Save ---
        supplier = str(record.get("GTK Supplier", "Unknown")).strip()
        platform = str(record.get("Platform", "Unknown")).strip()
        filename = f"Settlement Form_{supplier} {platform}.docx"
        out_path = self.output_dir / filename
        doc.save(str(out_path))
        return out_path


# ---------------------------------------------------------------------------
# Module-level helpers
# ---------------------------------------------------------------------------

def _build_replacements(
    record: dict,
    format_amount: Callable,
    amount_to_words: Callable,
    format_date: Callable,
) -> dict[str, str]:
    result: dict[str, str] = {}
    for placeholder, col in _PLACEHOLDER_MAP.items():
        if col == "__effective_date__":
            val = format_date(record.get("Effective Date", ""))
        elif col == "__actual_payment__":
            val = format_amount(record.get("Actual Payment", 0))
        elif col == "__capital_money__":
            val = amount_to_words(record.get("Actual Payment", 0))
        else:
            raw = record.get(col, "")
            val = "" if _is_nan(raw) else str(raw)
        result[placeholder] = val
    return result


def _is_nan(value: object) -> bool:
    """Return True for pandas NaN (float NaN != NaN)."""
    try:
        return float(value) != float(value)
    except (TypeError, ValueError):
        return False


def _replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    """Replace placeholders, handling cases where they span multiple runs."""
    if not any(ph in para.text for ph in replacements):
        return

    # Step 1: per-run replacement (preserves individual run formatting)
    for ph, val in replacements.items():
        for run in para.runs:
            if ph in run.text:
                run.text = run.text.replace(ph, val)

    # Step 2: if any placeholder is still present it must be split across runs;
    # consolidate all run text into the first run as a fallback.
    if any(ph in para.text for ph in replacements):
        full = para.text
        for ph, val in replacements.items():
            full = full.replace(ph, val)
        if para.runs:
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ""


def _apply_body_formatting(paragraphs: list) -> None:
    """Apply paragraph-level formatting rules to body paragraphs."""
    # Third paragraph (0-based index 2): bold + centred
    if len(paragraphs) > 2:
        p = paragraphs[2]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

    # "HP INC." and standalone "COMPANY": left-aligned bold
    for para in paragraphs:
        stripped = para.text.strip()
        if "HP INC." in stripped or stripped == "COMPANY":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True


def _apply_table_cell_formatting(para) -> None:
    """Right-align and bold the payment-total row inside tables."""
    if "Total Settlement Payment (in USD)" in para.text:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.bold = True


def _apply_global_font(doc: Document) -> None:
    """Set Times New Roman 10 pt on every run in the document."""

    def _fix(run) -> None:
        run.font.name = _FONT_NAME
        run.font.size = _FONT_SIZE

    for para in doc.paragraphs:
        for run in para.runs:
            _fix(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _fix(run)
